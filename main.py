from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.style import WD_BUILTIN_STYLE
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
from win32com import client
from re import findall
from os import walk,path
import time

def SetTitle2(file):
    word = client.DispatchEx("Word.Application")
    word.Visible = 0  # 设置应用可见
    word.DisplayAlerts = 0
    doc= word.Documents.Open(file)  # 使用微软office打开word
    Heading1 = doc.Styles(-2)
    Heading2 = doc.Styles(-3)
    FirstTitle1 = ['^一、.*?', '^二、.*?', '^三、.*?', '^四、.*?', '^五、.*?', '^六、.*?', '^七、.*?', '^八、.*?', '^九、.*?', '^十、.*?']
    FirstTitle2 = ['^第一章.*?', '^第二章.*?', '^第三章.*?', '^第四章.*?', '^第五章.*?', '^第六章.*?', '^第七章.*?', '^第八章.*?', '^第九章.*?',
                   '^第十章.*?']
    SecondTitle1 = ['^（一）.*?', '^（二）.*?', '^（三）.*?', '^（四）.*?', '^（五）.*?', '^（六）.*?', '^（七）.*?', '^（八）.*?', '^（九）.*?',
                    '^（十）.*?', ]
    SecondTitle2 = ['^第一节.*?', '^第二节.*?', '^第三节.*?', '^第四节.*?', '^第五节.*?', '^第六节.*?', '^第七节.*?', '^第八节.*?', '^第九节.*?',
                    '^第十节.*?']
    for para in doc.paragraphs:
        for first1,first2,second1,second2 in zip(FirstTitle1,FirstTitle2,SecondTitle1,SecondTitle2):
            #print(type(str(para)))
            if findall(first1,str(para)) != []:
                para.style = Heading1
            if findall(first2,str(para)) != []:
                para.style = Heading1
            if findall(second1,str(para)) != []:
                para.style = Heading2
            if findall(second2,str(para)) != []:
                para.style = Heading2
    NewName = file.split('.')[0]+'加标题.'+'docx'
    doc.SaveAs(NewName)
    doc.Close(SaveChanges=False)
    return NewName



def SetTitle(file):
    texts = Document(file)
    FirstTitle1 = ['^一、.*?','^二、.*?','^三、.*?','^四、.*?','^五、.*?','^六、.*?','^七、.*?','^八、.*?','^九、.*?','^十、.*?']
    FirstTitle2 = ['^第一章.*?','^第二章.*?','^第三章.*?','^第四章.*?','^第五章.*?','^第六章.*?','^第七章.*?','^第八章.*?','^第九章.*?','^第十章.*?']
    SecondTitle1 = ['^（一）.*?','^（二）.*?','^（三）.*?','^（四）.*?','^（五）.*?','^（六）.*?','^（七）.*?','^（八）.*?','^（九）.*?','^（十）.*?',]
    SecondTitle2 = ['^第一节.*?','^第二节.*?','^第三节.*?','^第四节.*?','^第五节.*?','^第六节.*?','^第七节.*?','^第八节.*?','^第九节.*?','^第十节.*?']
    for para in texts.paragraphs:
        for first1,first2,second1,second2 in zip(FirstTitle1,FirstTitle2,SecondTitle1,SecondTitle2):
            if findall(first1,para.text) != []:
                print("yes")
                para.style = 'Heading 1'
            if findall(first2,para.text) != []:
                para.style = 'Heading 1'
            if findall(second1,para.text) != []:
                para.style = 'Heading 2'
            if findall(second2,para.text) != []:
                para.style = 'Heading 2'
    for para in texts.paragraphs:
        print(para.text)
        print(para.style)
        pass
    NewName = file.split('.')[0]+'加标题.'+'docx'
    texts.save(NewName)
    return NewName

def AddContent(file):
    word = client.DispatchEx("Word.Application")
    word.Visible = 0  # 设置应用可见
    word.DisplayAlerts = 0
    doc= word.Documents.Open(file)  # 使用微软office打开word
    doc.Range(Start=0, End=0).InsertBreak()
    #doc.Paragraphs(1).Range.InsertBreak()  #第一段插入分页
    doc.Range(Start=0, End=0).InsertParagraphBefore()  # 在首行之前插入一行，用于插入目录
    FirstLineRange = doc.Paragraphs(1).Range  # 指向新插入的行
    FirstLineRange.Text = '目录'
    FirstLineRange.Font.Bold =True
    FirstLineRange.Font.Size = 20
    FirstLineRange.Font.Name = '仿宋'
    FirstLineRange.ParagraphFormat.Alignment =1

    FirstLineRange.InsertParagraphAfter()
    FirstLineRange.InsertParagraphAfter()
    SecondLineRange = doc.Paragraphs(2).Range

    doc.TablesOfContents.Add(Range=SecondLineRange,UseHeadingStyles=False,LowerHeadingLevel=2)  # 生成目录对象
    print(SecondLineRange.Text)

    NewName = file.split('.')[0]+'加目录.'+'docx'
    doc.SaveAs(NewName)
    doc.Close(SaveChanges=False)
    word.Quit()

def ContentAll(disc):
    filelist = []
    for root, dirs, files in walk(disc):
        for name in files:
            f = path.join(root, name)
            filelist.append(f)
    print(filelist)
    for file in filelist:
        AddContent(SetTitle2(file))

if __name__ == '__main__':
    #SetTitle(r'E:\Python\Projects\Bilibili\文本处理\Example\New - 副本.docx')
    #SetTitle2(r'E:\Python\Projects\Bilibili\文本处理\Example\New - 副本.docx')
    WenJianJia =  input('请拖入文件夹：')
    ContentAll(WenJianJia)